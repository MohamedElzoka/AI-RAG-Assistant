"""
rag/document_loader.py
======================
تحميل وتقسيم المستندات المختلفة لـ chunks صغيرة.

ليه بنقسم المستندات؟
- الـ LLMs عندها context window محدود
- الـ vector search أدق مع نصوص قصيرة
- الـ chunk_overlap بيضمن مانقطعش جملة في النص

أنواع الملفات المدعومة:
- .txt: ملفات نص عادي
- .md: Markdown files
- .pdf: PDF documents
- .docx: Word documents
"""

import re
from pathlib import Path
from typing import List, Optional, Dict
from dataclasses import dataclass

from llm.schemas import DocumentMetadata


@dataclass
class Document:
    """
    مستند واحد بعد تقسيمه.
    كل مستند هو chunk من النص الأصلي.
    """
    content: str                    # نص الـ chunk
    metadata: DocumentMetadata      # معلومات عن المصدر


class DocumentLoader:
    """
    تحميل المستندات وتقسيمها لـ chunks.

    الخطوات:
    1. قراءة الملف
    2. تنظيف النص
    3. تقسيمه لـ chunks بحجم محدد مع overlap
    """

    def __init__(
        self,
        chunk_size: int = 500,    # حجم كل chunk بالأحرف
        chunk_overlap: int = 50   # التداخل بين الـ chunks
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def load_file(self, file_path: str) -> List[Document]:
        """
        تحميل ملف وتحويله لقائمة من الـ documents.
        بيكتشف نوع الملف تلقائياً من الامتداد.
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"الملف مش موجود: {file_path}")

        # اختيار طريقة القراءة المناسبة حسب نوع الملف
        suffix = path.suffix.lower()

        if suffix == ".txt":
            text = self._load_txt(path)
        elif suffix == ".md":
            text = self._load_markdown(path)
        elif suffix == ".pdf":
            text = self._load_pdf(path)
        elif suffix == ".docx":
            text = self._load_docx(path)
        else:
            raise ValueError(f"نوع الملف غير مدعوم: {suffix}")

        # تنظيف النص
        text = self._clean_text(text)

        # تقسيم النص لـ chunks
        chunks = self._split_text(text)

        # تحويل الـ chunks لـ Document objects
        documents = []
        for i, chunk in enumerate(chunks):
            metadata = DocumentMetadata(
                source=str(path.name),
                doc_type=suffix.lstrip("."),
                chunk_index=i,
                total_chunks=len(chunks),
                char_count=len(chunk)
            )
            documents.append(Document(content=chunk, metadata=metadata))

        print(f"✅ تم تحميل: {path.name} → {len(documents)} chunk")
        return documents

    def load_directory(
        self,
        dir_path: str,
        extensions: Optional[List[str]] = None
    ) -> List[Document]:
        """
        تحميل كل الملفات في مجلد.

        Parameters:
            dir_path: مسار المجلد
            extensions: قائمة الامتدادات المطلوبة (default: كل الأنواع المدعومة)
        """
        if extensions is None:
            extensions = [".txt", ".md", ".pdf", ".docx"]

        path = Path(dir_path)
        all_documents = []

        for ext in extensions:
            for file_path in path.glob(f"**/*{ext}"):
                try:
                    docs = self.load_file(str(file_path))
                    all_documents.extend(docs)
                except Exception as e:
                    print(f"⚠️ فشل تحميل {file_path}: {e}")

        print(f"📚 إجمالي: {len(all_documents)} chunk من {dir_path}")
        return all_documents

    def load_text(self, text: str, source_name: str = "manual") -> List[Document]:
        """
        تحميل نص مباشرة (بدون ملف).
        مفيد للاختبار أو إضافة نصوص ديناميكية.
        """
        text = self._clean_text(text)
        chunks = self._split_text(text)

        documents = []
        for i, chunk in enumerate(chunks):
            metadata = DocumentMetadata(
                source=source_name,
                doc_type="text",
                chunk_index=i,
                total_chunks=len(chunks),
                char_count=len(chunk)
            )
            documents.append(Document(content=chunk, metadata=metadata))

        return documents

    # ==========================================
    # Private Methods - طرق التحميل الداخلية
    # ==========================================

    def _load_txt(self, path: Path) -> str:
        """قراءة ملف .txt"""
        return path.read_text(encoding="utf-8")

    def _load_markdown(self, path: Path) -> str:
        """قراءة ملف .md وإزالة رموز Markdown"""
        text = path.read_text(encoding="utf-8")
        # إزالة headers (#, ##, ###)
        text = re.sub(r'^#+\s+', '', text, flags=re.MULTILINE)
        # إزالة bold/italic
        text = re.sub(r'\*{1,2}([^*]+)\*{1,2}', r'\1', text)
        # إزالة links
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
        return text

    def _load_pdf(self, path: Path) -> str:
        """قراءة ملف PDF"""
        try:
            import pypdf
            text_parts = []
            with open(path, "rb") as f:
                reader = pypdf.PdfReader(f)
                for page in reader.pages:
                    text_parts.append(page.extract_text() or "")
            return "\n\n".join(text_parts)
        except ImportError:
            raise ImportError("pypdf مش مثبتة. شغّل: pip install pypdf")

    def _load_docx(self, path: Path) -> str:
        """قراءة ملف Word"""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(path)
            return "\n\n".join(para.text for para in doc.paragraphs if para.text)
        except ImportError:
            raise ImportError("python-docx مش مثبتة. شغّل: pip install python-docx")

    def _clean_text(self, text: str) -> str:
        """تنظيف النص من المسافات والأسطر الزيادة."""
        # تحويل tabs لمسافات
        text = text.replace("\t", " ")
        # إزالة أسطر فارغة متعددة متتالية
        text = re.sub(r'\n{3,}', '\n\n', text)
        # إزالة مسافات متعددة
        text = re.sub(r' {2,}', ' ', text)
        return text.strip()

    def _split_text(self, text: str) -> List[str]:
        """
        تقسيم النص لـ chunks.

        الخوارزمية:
        1. نحاول نقسم عند حدود الجمل أو الفقرات
        2. كل chunk بيبقى حجمه chunk_size تقريباً
        3. بيكون في overlap بين كل chunk والتالي
           عشان نحافظ على الـ context بين الـ chunks

        مثال مع chunk_size=20, overlap=5:
        النص: "ABCDEFGHIJKLMNOPQRSTUVWXYZ"
        Chunk 1: "ABCDEFGHIJKLMNOPQRST" (0-20)
        Chunk 2: "PQRSTUVWXYZ" (15-25) ← فيه overlap من P إلى T
        """
        if len(text) <= self.chunk_size:
            return [text] if text.strip() else []

        chunks = []
        start = 0

        while start < len(text):
            end = start + self.chunk_size

            if end >= len(text):
                # وصلنا لآخر النص
                chunks.append(text[start:].strip())
                break

            # نحاول نكسر عند نهاية جملة أو فقرة (مش في نص الكلمة)
            # بنبحث عن آخر نقطة/سطر جديد قبل الـ end
            split_pos = end
            for sep in ['\n\n', '\n', '. ', '! ', '? ']:
                last_sep = text.rfind(sep, start, end)
                if last_sep > start + self.chunk_size // 2:
                    split_pos = last_sep + len(sep)
                    break

            chunk = text[start:split_pos].strip()
            if chunk:
                chunks.append(chunk)

            # التالي يبدأ بـ overlap قبل نهاية الـ chunk الحالي
            start = split_pos - self.chunk_overlap

        return [c for c in chunks if c.strip()]
